"""Generate the nine attached input files for the Fairlead Street task."""
import sys, os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import site_data as s

from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = "/home/azureuser/geranium_tasks/task2_env/input_files"
os.makedirs(OUT, exist_ok=True)

GREEN = "1F4E3D"
HDR = Font(bold=True, color="FFFFFF", size=9)
THIN = Side(style="thin", color="BFBFBF")
BOX = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)

CAS = {"Tetrachloroethene": "127-18-4", "Trichloroethene": "79-01-6",
       "cis-1,2-Dichloroethene": "156-59-2", "Vinyl chloride": "75-01-4",
       "1,1,1-Trichloroethane": "71-55-6", "Methylene chloride": "75-09-2",
       "Arsenic": "7440-38-2", "Cadmium": "7440-43-9", "Chromium, total": "7440-47-3",
       "Lead": "7439-92-1", "Nickel": "7440-02-0"}
METHOD = {a: ("SW-846 8260D" if a in s.VOC else "SW-846 6020B") for a in s.VOC + s.MET}


def sheet(wb, title, headers, rows, widths=None, fmts=None, first=False, note=None):
    ws = wb.active if first else wb.create_sheet()
    ws.title = title
    r = 1
    if note:
        c = ws.cell(1, 1, note)
        c.font = Font(italic=True, size=8.5, color="595959")
        c.alignment = Alignment(wrap_text=True, vertical="top")
        ws.row_dimensions[1].height = 42
        r = 3
    for i, h in enumerate(headers, 1):
        c = ws.cell(r, i, h)
        c.font, c.fill, c.border = HDR, PatternFill("solid", fgColor=GREEN), BOX
        c.alignment = Alignment(wrap_text=True, vertical="center", horizontal="center")
    ws.row_dimensions[r].height = 30
    for j, row in enumerate(rows, r + 1):
        for i, v in enumerate(row, 1):
            c = ws.cell(j, i, v)
            c.border = BOX
            c.font = Font(size=9)
            if fmts and fmts.get(i):
                c.number_format = fmts[i]
    for i in range(1, len(headers) + 1):
        ws.column_dimensions[get_column_letter(i)].width = (widths or {}).get(i, 14)
    ws.freeze_panes = ws.cell(r + 1, 1)
    ws.sheet_view.showGridLines = False
    return ws


def doc_new(title, sub=None):
    d = Document()
    st = d.styles["Normal"]; st.font.name, st.font.size = "Calibri", Pt(10.5)
    p = d.add_paragraph(); run = p.add_run(title)
    run.bold, run.font.size, run.font.color.rgb = True, Pt(15), RGBColor(0x1F, 0x4E, 0x3D)
    if sub:
        q = d.add_paragraph(); r2 = q.add_run(sub)
        r2.italic, r2.font.size, r2.font.color.rgb = True, Pt(9.5), RGBColor(0x59, 0x59, 0x59)
    return d


def head(d, t, size=11.5):
    p = d.add_paragraph(); r = p.add_run(t)
    r.bold, r.font.size, r.font.color.rgb = True, Pt(size), RGBColor(0x1F, 0x4E, 0x3D)


# ============================================================ FILE 01 =======
def f01():
    d = doc_new("MARCHBANK ENVIRONMENTAL, INC.", "Internal assignment note")
    t = d.add_table(rows=5, cols=2); t.style = "Table Grid"
    for i, (k, v) in enumerate([
        ("To", "Project scientist, Site Assessment group"),
        ("From", "Dee Ferraro, PG, Principal"),
        ("Date", "March 9, 2026"),
        ("Project", f"{s.SITE}, {s.CITY} - Phase II ESA data evaluation (Proj. 26-0431)"),
        ("Client", f"{s.CLIENT}"),
    ]):
        t.rows[i].cells[0].width = Inches(0.8); t.rows[i].cells[1].width = Inches(5.9)
        t.rows[i].cells[0].text = ""; t.rows[i].cells[0].paragraphs[0].add_run(k).bold = True
        t.rows[i].cells[1].text = v
    d.add_paragraph()
    d.add_paragraph(
        "The Phase II field work wrapped up on February 25 and the last of the laboratory data landed "
        "Friday. Kestrel is under a purchase and sale agreement and their due diligence period expires "
        "at the end of this month, so they need our conclusions in time to decide whether to close, "
        "renegotiate or walk.")
    d.add_paragraph(
        "Marisa Quintanilla is their acquisitions lead and Ben Oyediran is outside counsel. Marisa wants "
        "a number she can take into a negotiation, not a range with caveats attached to every line. Ben "
        "wants to know whether anything we find falls inside or outside the indemnity. Read the PSA "
        "provisions before you write anything, because the answer to Ben's question changes the "
        "recommendation.")

    head(d, "What I need from you")
    for x in ["Whether the data are usable as reported, and where they are not, what that does to the "
              "conclusions. The laboratory did not perform validation and there are at least two things "
              "in the case narrative I want you to take a view on.",
              "Which constituents actually exceed the applicable criteria, by area and by exposure "
              "pathway. Be careful about which criteria apply here; do not default to the most "
              "conservative column in the table.",
              "Whether this site can be closed on a risk-based basis or whether it needs removal, and "
              "what that costs. Build the cost from the unit rates in our cost basis file, not from "
              "memory.",
              "How the cost sits against the escrow, and what Marisa should do about the gap.",
              "What is still unknown, what it would take to close those gaps, and whether that fits "
              "inside the due diligence clock."]:
        d.add_paragraph(x, style="List Number")

    head(d, "Things to watch")
    for x in ["Kestrel is building a distribution warehouse. They have already told us they will accept "
              "a deed restriction if it saves money, but the seller has to consent to one and the seller "
              "is keeping the parcel next door.",
              "The laboratory reported soil volatiles and soil metals in different units. I have had two "
              "reports come across my desk this year where that was missed.",
              "The disposal facility will not quote a rate until the soil is profiled. We have the "
              "profiling data.",
              "If you conclude we need more field work, tell me exactly how many days it adds. Marisa "
              "will ask and 'a few weeks' is not an answer she can use."]:
        d.add_paragraph(x, style="List Bullet")
    d.add_paragraph(
        "Write it as a memorandum to Kestrel over my signature. Assume Marisa is technical enough to "
        "follow a screening table but will not read an appendix.", style="Intense Quote")
    d.add_paragraph("- Dee")
    d.save(f"{OUT}/01_PM_Assignment_Note.docx")


# ============================================================ FILE 02 =======
def f02():
    wb = Workbook()
    rows = []
    for sid, b, aoc, top, bot in s.soil_samples():
        for a in s.VOC + s.MET:
            val, qual, rl = s.soil_result(sid, b, top, a)
            units = "ug/kg" if a in s.VOC else "mg/kg"
            mdl = round(rl / 3.0, 4)
            rows.append([f"L26-{7100 + len(rows)}", sid, "Soil", "2026-02-24" if b != "SB-12" else "2026-02-25",
                         "2026-03-03", METHOD[a], a, CAS[a], val, qual, rl, mdl, units, 1])
    sheet(wb, "Soil Results", ["Lab Sample ID", "Client Sample ID", "Matrix", "Date Collected",
                               "Date Analysed", "Method", "Analyte", "CAS", "Result", "Qualifier",
                               "RL", "MDL", "Units", "Dilution"], rows,
          widths={1: 13, 2: 15, 3: 8, 4: 12, 5: 12, 6: 15, 7: 24, 8: 12, 9: 11, 10: 10, 11: 9, 12: 9,
                  13: 8, 14: 8},
          fmts={9: '0.###', 11: '0.###', 12: '0.####'}, first=True,
          note=("Calderwood Analytical Laboratories, Report 26-1188. Client: Marchbank Environmental, "
                "Inc. Project 26-0431. Qualifier U indicates the analyte was not detected above the "
                "reporting limit; the value reported is the reporting limit. Data have not been "
                "independently validated. See the case narrative issued with this report."))

    tclp = [[f"{b}-{t:02d}{t+2:02d}", "SW-846 1311 / 6020B", "Lead", v, 0.10, "mg/L"]
            for (b, t), v in s.TCLP.items()]
    sheet(wb, "TCLP", ["Client Sample ID", "Method", "Analyte", "Result", "RL", "Units"], tclp,
          widths={1: 16, 2: 22, 3: 12, 4: 10, 5: 8, 6: 8}, fmts={4: '0.00'},
          note=("Toxicity characteristic leaching procedure performed at client request on the two soil "
                "samples with the highest total lead concentrations, for waste profiling."))
    wb.save(f"{OUT}/02_Soil_Analytical_Results.xlsx")


# ============================================================ FILE 03 =======
def f03():
    wb = Workbook()
    rows = []
    for w, aoc, _ in s.WELLS:
        for a in s.VOC + s.MET:
            val, qual, rl = s.gw_result(w, a)
            rows.append([f"L26-{7400 + len(rows)}", w, "Water", "2026-02-25",
                         "2026-03-11" if w == s.HOLDING_TIME_ISSUE and a in s.VOC else "2026-03-04",
                         METHOD[a], a, CAS[a], val, qual, rl, round(rl / 3.0, 3), "ug/L", 1])
    sheet(wb, "Groundwater Results", ["Lab Sample ID", "Client Sample ID", "Matrix", "Date Collected",
                                      "Date Analysed", "Method", "Analyte", "CAS", "Result",
                                      "Qualifier", "RL", "MDL", "Units", "Dilution"], rows,
          widths={1: 13, 2: 15, 3: 8, 4: 12, 5: 12, 6: 15, 7: 24, 8: 12, 9: 11, 10: 10, 11: 9,
                  12: 9, 13: 8, 14: 8},
          fmts={9: '0.###', 11: '0.###', 12: '0.###'}, first=True,
          note=("Calderwood Analytical Laboratories, Report 26-1191. Groundwater collected from "
                "temporary wells using low-flow methods. Qualifier U indicates not detected above the "
                "reporting limit. J indicates an estimated value. Data have not been independently "
                "validated. See the case narrative issued with this report."))

    gauge = [[w, note, dep, dep - 0.4] for (w, aoc, note), dep in
             zip(s.WELLS, [10.8, 10.4, 9.9, 12.1])]
    sheet(wb, "Water Level Gauging", ["Well", "Location", "Depth to Water (ft btoc)",
                                      "Depth to Water, corrected (ft bgs)"], gauge,
          widths={1: 10, 2: 42, 3: 20, 4: 24}, fmts={3: '0.0', 4: '0.0'},
          note="Gauged 2026-02-25 prior to purging. Inferred flow direction is east-southeast.")
    wb.save(f"{OUT}/03_Groundwater_Analytical_Results.xlsx")


# ============================================================ FILE 04 =======
def f04():
    wb = Workbook()
    rows = []
    for bid, aoc, td, water, note in s.BORINGS:
        rows.append([bid, s.AOC[aoc]["name"], td, water, note])
    sheet(wb, "Boring Summary", ["Boring", "Area", "Total Depth (ft bgs)",
                                 "First Water (ft bgs)", "Location Notes"], rows,
          widths={1: 10, 2: 52, 3: 16, 4: 16, 5: 46}, fmts={3: '0.0', 4: '0.0'}, first=True,
          note=("Marchbank Environmental field record, Phase II subsurface investigation, "
                "17 to 25 February 2026. Borings advanced by direct-push. Lithology across the site is "
                "3 to 5 ft of fill over silty sand to 9 to 12 ft, over a dense silt till."))

    samples = []
    for sid, b, aoc, top, bot in s.soil_samples():
        is_dup = sid == s.FIELD_DUP[2]
        samples.append([sid, b if not is_dup else s.FIELD_DUP[0], f"{top}-{bot}",
                        "Blind field duplicate" if is_dup else "Primary",
                        "Fill" if top < 4 else ("Silty sand" if top < 10 else "Silt till")])
    sheet(wb, "Soil Samples", ["Client Sample ID", "Boring", "Interval (ft bgs)", "Sample Type",
                               "Material"], samples,
          widths={1: 16, 2: 10, 3: 16, 4: 22, 5: 14},
          note=("Blind duplicates were submitted to the laboratory under a separate identifier. The "
                "parent boring and interval are recorded here."))

    feat = [[k, v["name"], v["dims"], v["area_sf"]] for k, v in s.AOC.items() if v["area_sf"]]
    sheet(wb, "Site Features", ["Area", "Feature", "Measured Dimensions", "Footprint (sq ft)"], feat,
          widths={1: 10, 2: 52, 3: 58, 4: 18}, fmts={4: '#,##0'},
          note=("Dimensions measured in the field against the 1994 as-built plan and the 1998 tank "
                "closure report."))
    wb.save(f"{OUT}/04_Field_Records_and_Site_Features.xlsx")


# ============================================================ FILE 05 =======
def f05():
    d = doc_new("CALDERWOOD ANALYTICAL LABORATORIES",
                "Case narrative for reports 26-1188 (soil) and 26-1191 (groundwater) | "
                "Client: Marchbank Environmental, Inc. | Project 26-0431 | Issued 2026-03-06")
    d.add_paragraph(
        "Samples were received in good condition, properly preserved and within temperature "
        "specification. All analyses were performed using the methods cited on the individual result "
        "pages. The following items are reported for the data user's consideration. Calderwood does not "
        "perform third-party data validation; qualification decisions beyond those shown on the result "
        "pages rest with the data user.")

    head(d, "1. Method blank, soil volatile organics, batch VB26-0451")
    d.add_paragraph(
        f"Methylene chloride was detected in the method blank associated with this batch at "
        f"{s.METHOD_BLANK['Methylene chloride']:g} ug/kg. Methylene chloride is a common laboratory "
        f"solvent and is present in the instrument background. All soil samples in report 26-1188 were "
        f"analysed in this batch. No other target analyte was detected in the method blank above the "
        f"reporting limit. Sample results have been reported as generated and have not been adjusted "
        f"for the blank result.")

    head(d, "2. Holding time, groundwater volatile organics, sample TW-03")
    d.add_paragraph(
        "The volatile organic aliquot for TW-03 was collected on 2026-02-25 and analysed on 2026-03-11, "
        "seventeen days after collection. The method holding time for preserved aqueous volatiles is "
        "fourteen days. The exceedance arose from an autosampler fault on 2026-03-04 that was not "
        "identified until the following week. Detected results for this sample are reported with a J "
        "qualifier and non-detected results with a UJ qualifier. Losses of volatile constituents over a "
        "three-day exceedance are in the low tens of percent for chlorinated ethenes under refrigerated "
        "storage. Re-collection was not possible within the reporting schedule.")

    head(d, "3. Field duplicate, soil, FD-01")
    d.add_paragraph(
        "The client submitted a blind field duplicate under identifier FD-01. Relative percent "
        "differences between FD-01 and its parent were within the laboratory's twenty percent control "
        "limit for chromium, cadmium and arsenic. The relative percent difference for lead was outside "
        "that limit. Field duplicate variability in heterogeneous fill material is commonly larger than "
        "laboratory duplicate variability and the laboratory control sample and matrix spike recoveries "
        "for lead in this batch were within limits, which indicates the variability is attributable to "
        "sample heterogeneity rather than to laboratory performance.")

    head(d, "4. Reporting limits, vinyl chloride in soil")
    d.add_paragraph(
        f"Vinyl chloride was not detected in any soil sample. The routine reporting limit for vinyl "
        f"chloride by SW-846 8260D in a soil matrix is {s.SOIL_RL['Vinyl chloride']:g} ug/kg. A lower "
        f"reporting limit of approximately 2 ug/kg is achievable by selected ion monitoring on request, "
        f"at additional cost and with a ten business day turnaround. This was not requested for the "
        f"present work.")

    head(d, "5. Surrogate and laboratory control sample recoveries")
    d.add_paragraph(
        "Surrogate recoveries were within acceptance limits for all samples. Laboratory control sample "
        "and laboratory control sample duplicate recoveries were within limits for all analytes in all "
        "batches. Matrix spike recovery for nickel in the batch containing SB-02 was 71 percent against "
        "a lower limit of 75 percent; the associated laboratory control sample recovered at 98 percent, "
        "indicating a matrix effect rather than a systematic bias.")

    head(d, "6. Continuing calibration")
    d.add_paragraph(
        "All continuing calibration verifications met method criteria. The initial calibration for the "
        "volatile organic instrument used for report 26-1188 was performed on 2026-02-19 with a mean "
        "relative response factor relative standard deviation of 8.4 percent.")
    d.save(f"{OUT}/05_Laboratory_Case_Narrative.docx")


# ============================================================ FILE 06 =======
def f06():
    wb = Workbook()
    rows = [[a, CAS[a], *s.SOIL_SL[a]] for a in s.VOC + s.MET]
    sheet(wb, "Soil", ["Constituent", "CAS", "Residential Direct Contact (mg/kg)",
                       "Industrial and Commercial Direct Contact (mg/kg)",
                       "Soil to Groundwater (mg/kg)"], rows,
          widths={1: 26, 2: 13, 3: 24, 4: 26, 5: 22}, fmts={3: '0.###', 4: '0.###', 5: '0.####'},
          first=True,
          note=("Ostrander Department of Environmental Quality, Risk-Based Screening Levels for Soil, "
                "Table 3-1, effective 2025-07-01. Direct contact levels are selected according to the "
                "documented use of the property. The soil to groundwater level protects the underlying "
                "aquifer and is applied at all properties irrespective of use classification."))

    grows = [[a, CAS[a], *[('' if v is None else v) for v in s.GW_SL[a]]] for a in s.VOC + s.MET]
    sheet(wb, "Groundwater", ["Constituent", "CAS", "GW-1 Potable Supply (ug/L)",
                              "GW-2 Vapour Migration to Buildings (ug/L)",
                              "GW-3 Discharge to Surface Water (ug/L)"], grows,
          widths={1: 26, 2: 13, 3: 22, 4: 26, 5: 24}, fmts={3: '0.###', 4: '0.###', 5: '0.###'},
          note=("Ostrander DEQ, Risk-Based Screening Levels for Groundwater, Table 4-2, effective "
                "2025-07-01. GW-1 applies where the groundwater is a current or potential source of "
                "potable supply, determined from the Department's aquifer classification mapping and "
                "the presence of supply wells. GW-2 applies where groundwater lies beneath, or within "
                "30 feet laterally of, an existing or planned occupied structure. GW-3 applies where "
                "groundwater discharges to a surface water body."))

    notes = [
        ["Use classification",
         "Industrial and commercial direct contact levels may be applied where the use is documented "
         "and where an activity and use limitation restricting residential occupancy is recorded "
         "against the deed."],
        ["Naturally occurring constituents",
         "A metal present at a concentration within the documented range for regional background soils "
         "is not considered to constitute a release. The Department's Brayton Lowland background study "
         "(2021) reports arsenic in native soils between 4 and 22 mg/kg."],
        ["Averaging",
         "Direct contact levels are applied to individual sample results within an exposure unit. "
         "Area-weighted averaging is not permitted for the construction worker receptor."],
        ["Reporting limits",
         "Where an analyte is not detected and the reporting limit achieved is above the applicable "
         "level, the result does not demonstrate compliance and re-analysis at a lower limit is "
         "required."],
        ["Activity and use limitations",
         "An activity and use limitation must be recorded before a closure certification is filed. "
         "Where the property is held subject to a purchase agreement, the party holding title at the "
         "time of recording must execute the instrument."],
    ]
    sheet(wb, "Application Notes", ["Topic", "Provision"], notes,
          widths={1: 28, 2: 108},
          note="Extracted from ODEQ Guidance Document RBSL-2025-01, sections 2 and 5.")
    wb.save(f"{OUT}/06_ODEQ_Screening_Levels.xlsx")


# ============================================================ FILE 07 =======
def f07():
    d = doc_new("PHASE I ENVIRONMENTAL SITE ASSESSMENT - FINDINGS SUMMARY",
                f"{s.SITE}, {s.CITY} | Prepared for {s.CLIENT} | Marchbank Environmental project "
                f"26-0431 | Issued 2026-01-22 | ASTM E1527-21")
    d.add_paragraph(
        "This summary carries forward the findings of the Phase I assessment that are relevant to the "
        "subsurface investigation. The full report is bound separately.")

    head(d, "Property and setting")
    for k, v in [
        ("Description", f"{s.ACRES} acres improved with a 48,000 square foot single-storey masonry "
                        "manufacturing building (Building A), a paved yard and a former tank pad."),
        ("Zoning", "M-2 Light Industrial. Residential occupancy is not a permitted use under the "
                   "Brayton zoning ordinance."),
        ("Buyer's intended use", "Demolition of Building A and construction of a 62,000 square foot "
                                 "distribution warehouse on slab, with the building footprint covering "
                                 "the east yard and the north yard."),
        ("Water supply", "The property and all surrounding properties are served by the Brayton "
                         "municipal water system, sourced from the Ferris Reservoir surface water "
                         "intake 7 miles north-west. A records search of the Ostrander DEQ well "
                         "registry identified no private supply wells within one mile."),
        ("Aquifer classification", "The surficial aquifer beneath the property is mapped by ODEQ as "
                                   "Class II-B, not a current or potential source of potable supply, "
                                   "on the basis of documented yields below 5 gallons per minute and "
                                   "the availability of municipal service."),
        ("Surface water", "The nearest surface water is Cobb Brook, 2,100 feet east and "
                          "hydraulically downgradient."),
        ("Groundwater", "First encountered at 10 to 12 feet below grade in nearby borings; inferred "
                        "flow east-southeast."),
    ]:
        p = d.add_paragraph(); p.add_run(f"{k}. ").bold = True; p.add_run(v)

    head(d, "Site history")
    d.add_paragraph(
        "The property was developed in 1961 for Fairlead Metal Products, which operated a metal "
        "fabrication and finishing plant until 2009. Operations included stamping, welding, vapour "
        "degreasing and an electroplating line producing chromium and cadmium finishes. The plant "
        "closed in 2009 and the building has been vacant since 2014 apart from intermittent storage "
        "use by the current owner.")

    head(d, "Recognised environmental conditions")
    for t, txt in [
        ("REC-1, former electroplating line",
         "A plating line with chromic acid, cadmium cyanide and nickel baths occupied the north-central "
         "portion of Building A. Photographs from a 1996 fire inspection show open rinse tank pits in "
         "the floor slab. Bath contents were reportedly removed in 2010 but no disposal documentation "
         "was located."),
        ("REC-2, former degreaser aboveground tank",
         "A 1,500 gallon aboveground tank on a concrete pad in the east yard supplied a vapour "
         "degreaser. Interviews with a former maintenance supervisor indicate the tank held "
         "tetrachloroethene and that the pad was not curbed. The tank was removed in 2011."),
        ("REC-3, former floor-drain dry well",
         "The 1994 as-built plan shows Building A floor drains discharging to a dry well in the north "
         "yard. There is no record of a connection to the municipal sewer before 1988. The structure "
         "was located in the field and remains in place."),
        ("REC-4, former underground storage tank",
         "A 4,000 gallon heating oil underground storage tank was removed from the west yard in 1998. "
         "The ODEQ tank closure file records a release, a subsequent excavation of approximately 180 "
         "tons of soil, and a no further action determination issued 1999-04-14 under the standards "
         "then in force. The file notes that post-excavation soil concentrations were left in place at "
         "levels above current criteria."),
    ]:
        head(d, t, size=10.5)
        d.add_paragraph(txt)

    head(d, "Regional background")
    d.add_paragraph(
        "The ODEQ Brayton Lowland background soil study (2021) characterises native soils across the "
        "surrounding area. Arsenic in native soils ranges from 4 to 22 mg/kg with a mean of 11 mg/kg, "
        "reflecting the glaciolacustrine parent material. Lead, chromium, cadmium and nickel in native "
        "soils are below the corresponding residential direct contact levels throughout the study area.")
    d.save(f"{OUT}/07_Phase_I_Findings_Summary.docx")


# ============================================================ FILE 08 =======
def f08():
    wb = Workbook()
    rows = [[n, b, r, u] for n, b, r, u in s.UNIT_COSTS]
    sheet(wb, "Unit Costs", ["Item", "Basis", "Rate (USD)", "Unit"], rows,
          widths={1: 62, 2: 30, 3: 14, 4: 8}, fmts={3: '#,##0.00'}, first=True,
          note=("Marchbank Environmental estimating basis, revised 2026-01. Rates are derived from "
                "three remediation projects bid in the Brayton and Coldharbour markets during 2025 and "
                "carry a stated accuracy of minus 15 to plus 30 percent. Rates exclude contingency."))

    conv = [
        ["In-situ soil density, silty sand and fill", s.SOIL_DENSITY, "tons per in-situ cubic yard",
         "Weighted from three 2025 projects in comparable material."],
        ["Cubic feet per cubic yard", 27, "cf/cy", ""],
        ["Contingency, pre-design estimate", s.CONTINGENCY, "fraction",
         "Applied to the subtotal for estimates prepared before a remedial design."],
        ["Confirmation sample sets assumed", s.CONFIRMATION_SETS, "sets",
         "Four sidewall and one floor set per excavation area, plus stockpile characterisation."],
        ["Monitoring rounds assumed", s.MONITORING_ROUNDS, "rounds",
         "Four consecutive quarters is the minimum ODEQ will accept to support a closure filing."],
    ]
    sheet(wb, "Conversions and Assumptions", ["Parameter", "Value", "Unit", "Note"], conv,
          widths={1: 42, 2: 12, 3: 28, 4: 66})

    disp = [
        ["Ridgeline Regional Landfill, Coldharbour", "Non-hazardous industrial special waste",
         "Accepts soil profiled below RCRA toxicity characteristic limits. Profile approval requires "
         "TCLP results for the characteristic metals. 30 mile haul."],
        ["Tallmadge Treatment and Recovery", "RCRA hazardous soil, D008",
         "Required only where the toxicity characteristic leaching procedure result equals or exceeds "
         "5.0 mg/L for lead. 145 mile haul, manifested."],
    ]
    sheet(wb, "Disposal Facilities", ["Facility", "Waste Classification", "Acceptance Conditions"],
          disp, widths={1: 40, 2: 38, 3: 88})
    wb.save(f"{OUT}/08_Remedial_Cost_Basis.xlsx")


# ============================================================ FILE 09 =======
def f09():
    d = doc_new("PURCHASE AND SALE AGREEMENT - ENVIRONMENTAL PROVISIONS",
                f"Extract prepared by Oyediran Vance LLP for {s.CLIENT} | 2026-03-06 | "
                f"Privileged and confidential")
    d.add_paragraph(
        "The following provisions of the Purchase and Sale Agreement dated 2026-01-08 between Fairlead "
        "Holdings LLC as seller and Kestrel Industrial Partners LLC as buyer bear on the environmental "
        "assessment. Defined terms carry the meanings given in the Agreement.")

    for t, body in [
        ("Section 4.2, Due Diligence Period",
         "The Due Diligence Period commences on the Effective Date and expires at 5:00 p.m. Eastern "
         "time on March 31, 2026. Buyer may terminate this Agreement for any reason or no reason by "
         "written notice delivered before expiry, whereupon the Deposit shall be returned in full."),
        ("Section 4.3, Extension",
         "Buyer may extend the Due Diligence Period on one occasion by a period not exceeding sixty "
         "(60) days by written notice delivered before expiry, accompanied by an additional deposit of "
         "$75,000, which shall become non-refundable upon delivery but shall be credited against the "
         "Purchase Price at Closing. Seller's consent is not required."),
        ("Section 9.1, Environmental Escrow",
         "At Closing, Seller shall deposit into escrow the sum of Three Hundred Fifty Thousand Dollars "
         "($350,000) to be applied against the cost of Remedial Work. Seller's aggregate liability for "
         "Remedial Work shall not exceed the escrowed amount, and Buyer shall bear all cost in excess "
         "of that sum."),
        ("Section 9.2, Indemnity and Known Conditions",
         "Seller shall indemnify Buyer against Environmental Losses arising from Pre-Closing Releases, "
         "provided that this indemnity shall not extend to any Known Condition. 'Known Condition' means "
         "any condition identified or described in the Phase I Environmental Site Assessment prepared "
         "by Marchbank Environmental, Inc. and dated January 22, 2026, or in any governmental file "
         "referenced therein."),
        ("Section 9.4, Institutional Controls",
         "Buyer shall not record, and shall not permit to be recorded, any activity and use limitation, "
         "environmental land use restriction, notice of contamination or comparable instrument against "
         "the Property without Seller's prior written consent, which consent shall not be unreasonably "
         "withheld or delayed. Seller has disclosed that it will retain the adjoining parcel at 1462 "
         "Fairlead Street following Closing."),
        ("Section 9.5, Remedial Work",
         "'Remedial Work' means any investigation, removal, treatment, monitoring, institutional "
         "control or other response action required by an Environmental Authority, or reasonably "
         "necessary to obtain a closure determination, in respect of a Pre-Closing Release."),
        ("Section 11.6, Closing",
         "Closing shall occur no later than thirty (30) days following expiry of the Due Diligence "
         "Period, as the same may be extended under Section 4.3."),
    ]:
        head(d, t, size=10.5)
        d.add_paragraph(body)

    head(d, "Counsel's note")
    d.add_paragraph(
        "Two points for the consultant. First, Section 9.2 turns on what the Phase I identified, not on "
        "what caused the impact; if a condition is described in that report or in a file it cites, the "
        "indemnity does not reach it and the cost sits with Buyer above the escrow. Second, Section 9.4 "
        "means any remedy that depends on a deed restriction is not within Buyer's unilateral control. "
        "If your recommendation relies on one, say so plainly and tell us what the alternative costs, "
        "because that is the number I will be negotiating against. - B. Oyediran")
    d.save(f"{OUT}/09_PSA_Environmental_Provisions.docx")


for fn in (f01, f02, f03, f04, f05, f06, f07, f08, f09):
    fn()
    print("built", fn.__name__)
print("\n".join(sorted(os.listdir(OUT))))
